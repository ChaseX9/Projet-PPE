"""
Script de génération de l'étude de faisabilité au format Word.
Projet : Robo-Advisor (CapInvest) — Version Mars 2026
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "Etude_Faisabilite_RoboAdvisor.docx")

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    colors = {1: RGBColor(0x1A,0x2E,0x5A), 2: RGBColor(0x1F,0x6E,0xB4), 3: RGBColor(0x2E,0x86,0xAB)}
    sizes  = {1: 16, 2: 13, 3: 12}
    run.font.color.rgb = colors.get(level, RGBColor(0x33,0x33,0x33))
    run.font.size = Pt(sizes.get(level, 11))
    return p

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x33,0x33,0x33)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33,0x33,0x33)
    return p

def make_table(doc, headers, rows, header_color='1F6EB4', alt_color='EFF4FB'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, header_color)
        r = c.paragraphs[0].runs[0]
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            if ri % 2 == 0:
                set_cell_bg(cells[ci], alt_color)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    return t

def color_last_col(table, col_idx, color_map):
    """Color cells in the last column based on keyword presence."""
    for row in table.rows[1:]:
        cell = row.cells[col_idx]
        text = cell.text
        bg = 'FFFFFF'
        for kw, col in color_map.items():
            if kw in text:
                bg = col
                break
        set_cell_bg(cell, bg)

# ─── Document ────────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3);  section.right_margin  = Cm(3)

# ══════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ÉTUDE DE FAISABILITÉ")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1A,0x2E,0x5A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run("Projet PPE — Robo-Advisor « CapInvest »")
r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x1F,0x6E,0xB4)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ["Date : Mars 2026\n", "Équipe : 4 étudiants\n", "Méthode : TELOS\n"]:
    run = meta.add_run(line)
    run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x66,0x66,0x66)

doc.add_page_break()

# ══════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════
add_heading(doc, "1. Introduction", 1)

add_heading(doc, "1.1 Contexte du projet", 2)
add_para(doc,
    "Notre solution consiste en la conception et le développement d'un Robo-Advisor nommé « CapInvest ». "
    "Ce projet collectif, mené par une équipe de 4 étudiants, vise à démocratiser la gestion de "
    "portefeuille financier en offrant à des utilisateurs non-experts et pris par le temps, "
    "un outil automatisé, transparent et conforme aux réglementations européennes (MiFID II)."
)
add_para(doc,
    "L'application est une plateforme web full-stack développée en Python, proposant : un "
    "questionnaire de profilage intelligent (Safety-First), une optimisation mathématique de "
    "portefeuille (Markowitz et Black-Litterman), un explorateur d'univers d'actifs, une section "
    "éducative (CapInvest Academy), un export PDF et une gestion des données utilisateurs conforme au RGPD."
)

add_heading(doc, "1.2 Objectifs de l'étude", 2)
for obj in [
    "Vérifier que le projet est réalisable et viable dans le contexte d'un projet pluridisciplinaire en équipe.",
    "Identifier les freins et risques potentiels rencontrés ou anticipés.",
    "Justifier les choix technologiques avec des éléments factuels issus du code produit.",
    "Formuler des recommandations claires pour la suite du projet.",
]:
    add_bullet(doc, obj)

add_heading(doc, "1.3 Périmètre de l'étude", 2)
add_para(doc, "L'étude porte sur l'ensemble des composants réalisés :")
for scope in [
    "Backend API REST (Python / FastAPI) — 5 routers, 30+ endpoints",
    "Base de données locale (SQLite / SQLAlchemy) — 8 modèles ORM",
    "Moteur d'optimisation financière (Markowitz, Black-Litterman, CVXPY)",
    "Système d'authentification complet (JWT, bcrypt, vérification e-mail, reset MDP)",
    "Interface utilisateur (HTML / CSS / Vanilla JS) — 16 pages",
    "Explorateur d'univers d'actifs + Export PDF côté client",
    "Module éducatif : CapInvest Academy (leçons, quiz, XP, streaks)",
    "Conformité réglementaire (RGPD, MiFID II)",
    "Configuration de déploiement cloud (Render)",
]:
    add_bullet(doc, scope)

add_para(doc,
    "\nSont exclus du périmètre : l'intégration d'un vrai courtier financier et "
    "les tests de charge à grande échelle.", italic=True)
doc.add_paragraph()

# ══════════════════════════════════════════════
# 2. MÉTHODOLOGIE
# ══════════════════════════════════════════════
add_heading(doc, "2. Méthodologie", 1)
add_para(doc,
    "L'analyse s'appuie sur la méthode TELOS. Les données proviennent exclusivement du "
    "travail réellement accompli dans le code source :"
)
for src in [
    "Code source (dossiers src/, scripts/, templates/, static/)",
    "Documentation du projet (README.md, PROJECT_HANDOVER.md)",
    "Requirements techniques fixés dans requirements.txt (30 dépendances)",
    "Documentation officielle FastAPI, SQLAlchemy, PyPortfolioOpt, CVXPY",
    "Référentiel MiFID II — Directive 2014/65/UE",
    "RGPD — Règlement UE 2016/679",
]:
    add_bullet(doc, src)
doc.add_paragraph()

# ══════════════════════════════════════════════
# 3. ANALYSE TELOS
# ══════════════════════════════════════════════
add_heading(doc, "3. Analyse TELOS", 1)

# ── 3.1 Technique ──────────────────────────────
add_heading(doc, "3.1 Technique (T)", 2)
add_para(doc,
    "La dimension technique évalue si les technologies, compétences et ressources matérielles "
    "disponibles sont suffisantes pour réaliser le projet."
)

add_heading(doc, "Technologies effectivement utilisées", 3)
tech_rows = [
    ("FastAPI 0.129", "Framework API REST", "5 routers montés dans main.py (auth, portfolio, training, RGPD, core). 30+ endpoints.", "✅ Maîtrisé"),
    ("SQLite + SQLAlchemy 2.0", "SGBD local + ORM", "8 modèles : User, SavedProfile, SavedPortfolio, Module, Lesson, Question, UserLessonProgress, AuthToken.", "✅ Fonctionnel"),
    ("yfinance 1.1 + Tenacity", "Données financières", "479 actifs (405 Actions + 74 ETFs — US, Europe, Global, Émergents). Cache 48h (universe.csv). Tenacity gère les retries automatiques sur yfinance.", "✅ Résilient"),
    ("PyPortfolioOpt + CVXPY + scipy", "Optimisation", "Markowitz Mean-Variance (optimizer.py). CVXPY/scipy comme solveurs avancés. Contraintes 2–10 actifs, concentration 30–50%.", "✅ Opérationnel"),
    ("Black-Litterman", "Optimisation expert", "Réservé profils « Expert ». Verrouillage automatique si profil novice (gating MiFID II).", "✅ Implémenté"),
    ("JWT + bcrypt", "Auth sécurisée", "Tokens JWT (python-jose), mots de passe hachés (passlib/bcrypt), validation force MDP (majuscule, caract. spécial).", "✅ Sécurisé"),
    ("Vérif. e-mail + Reset MDP", "Auth avancée", "AuthToken (type verify/reset, expiration 24h/1h). Emails envoyés via service externe. Pages verify_status.html, reset_password.html.", "✅ Opérationnel"),
    ("Plotly", "Visualisation", "Graphiques interactifs dans l'explorateur d'univers d'actifs (explorer.html).", "✅ Intégré"),
    ("html2pdf + jsPDF", "Export PDF", "Export PDF côté client depuis le navigateur (aucun traitement serveur requis).", "✅ Fonctionnel"),
    ("HTML/CSS/Vanilla JS", "Frontend", "16 templates HTML. JS structuré : api.js, ui.js, smart-suggestions.js, cookie-consent.js, tour.js.", "✅ Complet"),
    ("Render (render.yaml)", "Déploiement cloud", "Configuration Render prête : runtime Python, startCommand uvicorn, PYTHON_VERSION=3.10.", "✅ Configuré"),
]

tech_table = make_table(doc,
    ["Technologie", "Rôle", "Usage réel dans le projet", "Statut"],
    tech_rows, header_color='1F6EB4'
)
status_colors = {"✅": "D4EDDA", "⚠️": "FFF3CD", "🔄": "D1ECF1", "❌": "F8D7DA"}
for row in tech_table.rows[1:]:
    cell = row.cells[3]
    for kw, col in status_colors.items():
        if kw in cell.text:
            set_cell_bg(cell, col)
            break
    # reset alt-row on status col only
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(9)

doc.add_paragraph()
add_para(doc,
    "Le projet a été développé sur machines locales (Mac/Windows) sans infrastructure cloud "
    "obligatoire. Toutes les bibliothèques sont open-source et disponibles via pip. "
    "La complexité technique réside dans les modèles financiers (Markowitz, Black-Litterman, CVXPY) "
    "et dans la logique MiFID II (Safety-First, gating, cohérence).", italic=True
)
add_para(doc, "\nConclusion Technique : ", bold=True)
add_para(doc,
    "Le projet est techniquement faisable et démontre une stack complète et maîtrisée. "
    "Le seul risque résiduel concerne la disponibilité de l'API yfinance (externe, non garantie), "
    "partiellement mitigé par la bibliothèque Tenacity et le cache local."
)
doc.add_paragraph()

# ── 3.2 Économique ─────────────────────────────
add_heading(doc, "3.2 Économique (E)", 2)
add_para(doc, "Évaluation de la soutenabilité financière du projet.")

econ_rows = [
    ("Hébergement local", "0 €", "Application run en localhost:8000. Aucun serveur externe requis en phase dev."),
    ("Hébergement cloud (Render)", "0 € (tier gratuit)", "render.yaml configuré pour déploiement sur Render Free Tier."),
    ("Base de données", "0 €", "SQLite : inclus dans Python, aucune licence requise."),
    ("Données financières", "0 €", "yfinance (API gratuite Yahoo Finance). 154 actifs sans abonnement."),
    ("Bibliothèques Python", "0 €", "30 dépendances open-source : FastAPI, SQLAlchemy, CVXPY, Plotly, etc."),
    ("Service e-mail", "0 € (dev)", "Service d'envoi e-mail (vérif. compte, reset MDP). Gratuit en phase académique."),
    ("Outils de développement", "0 €", "VS Code (gratuit), Git (gratuit), terminal Mac/Windows."),
    ("Matériel informatique", "Existant", "Machines personnelles de l'équipe — aucun achat requis."),
    ("TOTAL", "0 €", "Projet 100% gratuit dans sa phase de développement et de démo."),
]

econ_table = make_table(doc, ["Poste de coût", "Montant", "Justification"], econ_rows, header_color='1F6EB4')
last_row = econ_table.rows[-1].cells
for c in last_row:
    set_cell_bg(c, 'D4EDDA')
    for para in c.paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(9)

doc.add_paragraph()
add_para(doc,
    "En production réelle (post-PPE), des coûts apparaîtraient : abonnement API financière "
    "professionnelle (Bloomberg, Refinitiv), service e-mail payant, et éventuellement un agrément AMF. "
    "Ces coûts dépassent le cadre académique actuel.", italic=True
)
add_para(doc, "\nConclusion Économique : ", bold=True)
add_para(doc,
    "Le projet est économiquement viable dans son périmètre académique grâce à l'utilisation "
    "exclusive de technologies gratuites et open-source, y compris pour le déploiement cloud."
)
doc.add_paragraph()

# ── 3.3 Légal ──────────────────────────────────
add_heading(doc, "3.3 Légal (L)", 2)
add_para(doc, "Examen des contraintes réglementaires applicables au projet.")

add_heading(doc, "Réglementation MiFID II", 3)
add_para(doc,
    "MiFID II impose l'évaluation du profil de risque des clients avant toute recommandation "
    "d'investissement. Le projet intègre ces exigences à plusieurs niveaux :"
)
mifid_items = [
    "Questionnaire de profilage obligatoire : objectifs, horizon, tolérance au risque, capacité de perte, niveau de connaissance (src/profile/suggestion_rules.py).",
    "Logique Safety-First : la capacité de perte prime toujours sur l'objectif de performance (Tier 1 > Tier 2). Un utilisateur sans capacité de perte est systématiquement orienté vers un profil Prudent.",
    "Contrôles de cohérence MiFID II (profile_engine.py) : détection des incohérences (ex. tolérance haute + capacité nulle) avec message pédagogique.",
    "Score de confiance (confidence_score) : 0.95 si cohérence parfaite, 0.65 si flags détectés.",
    "Verrouillage expert : Black-Litterman réservé aux profils « Expert ». Passage en novice = désactivation automatique.",
    "Transparence : chaque actif du portefeuille est accompagné d'une explication en français (Core, Diversification, Fiabilité) via explainer.py.",
    "Exclusion d'actifs par l'utilisateur (excluded_tickers dans SavedPortfolio) : respect de l'autonomie du client.",
    "Aucune boîte noire : toutes les décisions d'allocation sont traçables et documentées.",
]
for item in mifid_items:
    add_bullet(doc, item)

add_heading(doc, "Réglementation RGPD", 3)
add_para(doc,
    "Le RGPD (UE 2016/679) impose des obligations strictes sur la collecte, le traitement "
    "et la suppression des données personnelles."
)
rgpd_items = [
    "Modèle User avec champs de consentement explicites : privacy_policy_accepted_at, cookies_consent, marketing_consent (src/database/models.py).",
    "Activation du compte uniquement après vérification e-mail (AuthToken type='verify', expiration 24h) — privacy by design.",
    "Reset de mot de passe sécurisé (AuthToken type='reset', expiration 1h). Tokens à usage unique, invalidés après utilisation.",
    "Endpoints RGPD dédiés (gdpr_endpoints.py) : export des données personnelles, suppression du compte.",
    "Page de gestion des données (gestion_donnees.html), politique de confidentialité (politique_confidentialite.html), mentions légales (mentions_legales.html).",
    "Gestion du consentement cookies côté client (cookie-consent.js).",
    "Mots de passe hachés bcrypt — aucun mot de passe en clair stocké.",
    "Tokens JWT avec expiration pour l'authentification (python-jose).",
    "Validation de force du mot de passe (majuscule + caractère spécial obligatoires).",
]
for item in rgpd_items:
    add_bullet(doc, item)

add_para(doc, "\nConclusion Légale : ", bold=True)
add_para(doc,
    "Le projet démontre une prise en compte très sérieuse des contraintes réglementaires. "
    "MiFID II et RGPD ont été intégrés by design, dès la conception. L'architecture respecte "
    "les principes fondamentaux qui seraient exigés en production réelle."
)
doc.add_paragraph()

# ── 3.4 Organisationnel ────────────────────────
add_heading(doc, "3.4 Organisationnel (O)", 2)
add_para(doc,
    "La dimension organisationnelle évalue les ressources humaines, la logistique et la gestion "
    "du projet au sein de l'équipe."
)

add_heading(doc, "Équipe et répartition du travail", 3)
add_para(doc,
    "Le projet a été conduit par une équipe de 4 étudiants, avec une organisation par domaines "
    "de compétences et un développement itératif structuré en phases successives :"
)
org_items = [
    "Développement itératif par phases : le projet a été découpé en 12 phases (Phase 1 : modèle mathématique → Phase 12 : déploiement cloud), permettant des livraisons progressives et testables.",
    "Documentation interne : README.md maintenu en continu comme documentation principale du projet.",
    "Scripts de maintenance : 24 scripts dans scripts/ (seeding, migration, debug, refresh univers) structurent les opérations récurrentes.",
    "Tests automatisés : dossier tests/ avec pytest et httpx.",
    "Gestion multi-plateforme : scripts de lancement pour Mac (lancer_app_mac.sh) et Windows (lancer_app_win.bat).",
    "Versioning : projet géré sous Git (.gitignore configuré).",
]
for item in org_items:
    add_bullet(doc, item)

add_heading(doc, "Architecture technique bien structurée", 3)
arch_items = [
    "src/api/ : 5 routers (core, auth, portfolio, training, RGPD) — 30+ endpoints",
    "src/portfolio/ : optimizer, recommender, explainer, filters, allocation",
    "src/profile/ : suggestion_rules, suggestion_engine, profile_engine, profile_builder, schemas",
    "src/database/ : 8 modèles SQLAlchemy + seeding automatique",
    "src/auth/ + src/utils/ : auth, dépendances, email_service, config",
    "templates/ : 16 pages HTML",
    "static/ : CSS + 7 fichiers JS (api.js, ui.js, smart-suggestions.js, cookie-consent.js, html2pdf, jsPDF, tour.js)",
    "scripts/ : 24 scripts utilitaires",
]
for item in arch_items:
    add_bullet(doc, item)

add_para(doc, "\nConclusion Organisationnelle : ", bold=True)
add_para(doc,
    "La structure du projet est solide et bien organisée. La décomposition en 12 phases claires, "
    "la documentation soignée, la séparation stricte des responsabilités et la collaboration en équipe "
    "de 4 témoignent d'une rigueur organisationnelle adaptée à un projet de cette envergure."
)
doc.add_paragraph()

# ── 3.5 Scheduling ─────────────────────────────
add_heading(doc, "3.5 Scheduling — Calendrier (S)", 2)
add_para(doc,
    "Évaluation du respect du calendrier et de la priorisation des phases."
)

sched_rows = [
    ("Phase 1–2", "Déc. 2025", "Moteur mathématique Markowitz + plateforme web de base (FastAPI, SQLite, JWT)", "✅ Livré"),
    ("Phase 3–4", "Janv. 2026", "Black-Litterman + raffinement UX (explications, profil utilisateur)", "✅ Livré"),
    ("Phase 5–6", "Janv. 2026", "Contraintes réalistes (2–10 actifs), cycle de vie portefeuille (accepter/refuser), exclusion actifs", "✅ Livré"),
    ("Phase 7–9", "Fév. 2026", "Transparence, moteur suggestions temps réel, Safety-First, conformité MiFID II complète", "✅ Livré"),
    ("Phase 10–11", "Fév. 2026", "Expansion univers (479 actifs), classification dynamique des profils, verrouillage Black-Litterman, RGPD complet", "✅ Livré"),
    ("Academy", "Fév. 2026", "Modules pédagogiques, leçons, quiz (QCM / vrai-faux / numérique), XP, streaks, gamification", "✅ Livré"),
    ("Phase 12", "Mars 2026", "Vérification e-mail, reset MDP, explorateur d'univers d'actifs (explorer.html), export PDF, Plotly (graphiques interactifs), CVXPY, configuration Render", "✅ Livré"),
]

sched_table = make_table(doc, ["Phase", "Période", "Contenu livré", "Statut"],
    sched_rows, header_color='1F6EB4')
for row in sched_table.rows[1:]:
    set_cell_bg(row.cells[3], 'D4EDDA')

doc.add_paragraph()
add_para(doc, "\nConclusion Scheduling : ", bold=True)
add_para(doc,
    "Le calendrier a été respecté. Les 12 phases de développement ont été réalisées en "
    "environ 3 mois (décembre 2025 – mars 2026), avec une progression cohérente du MVP vers "
    "un produit complet et deployable. La principale contrainte du PPE (délai de livraison) "
    "a été anticipée par un découpage en phases livrables indépendamment."
)
doc.add_paragraph()

# ══════════════════════════════════════════════
# 4. CONCLUSIONS ET RECOMMANDATIONS
# ══════════════════════════════════════════════
add_heading(doc, "4. Conclusions et recommandations", 1)

add_heading(doc, "4.1 Synthèse TELOS", 2)
synth_rows = [
    ("T – Technique", "Stack complète maîtrisée. Markowitz + Black-Litterman + CVXPY. Tenacity pour la résilience yfinance.", "✅ Favorable"),
    ("E – Économique", "Coût zéro en phase académique. Déploiement cloud gratuit (Render). Stack 100% open-source.", "✅ Favorable"),
    ("L – Légal", "MiFID II et RGPD intégrés by design. Vérif. e-mail, reset MDP, consentements, export/suppression données.", "✅ Favorable"),
    ("O – Organisationnel", "Équipe de 4. 12 phases bien découpées, architecture claire, documentation rigoureuse.", "✅ Favorable"),
    ("S – Scheduling", "12 phases livrées en 3 mois. Calendrier respecté, progression maîtrisée.", "✅ Favorable"),
]
synth_table = make_table(doc, ["Dimension", "Résumé", "Verdict global"],
    synth_rows, header_color='1A2E5A')
for row in synth_table.rows[1:]:
    cell = row.cells[2]
    if "✅" in cell.text:
        set_cell_bg(cell, 'D4EDDA')
    elif "⚠️" in cell.text:
        set_cell_bg(cell, 'FFF3CD')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, "4.2 Faisabilité globale", 2)
add_para(doc,
    "Le projet CapInvest Robo-Advisor est globalement faisable et a été réalisé avec succès. "
    "Les 5 dimensions TELOS sont toutes favorables. Sur le plan technique, économique, légal, "
    "organisationnel et temporel, le projet répond aux critères d'un développement rigoureux et viable."
)

add_heading(doc, "4.3 Risques identifiés et mitigation", 2)
risk_rows = [
    ("Dépendance à Yahoo Finance", "Moyen", "API gratuite non garantie (rate limits, données manquantes)", "Cache 48h (universe.csv) + Tenacity pour retries automatiques + script refresh manuel"),
    ("Service e-mail externe", "Faible", "L'activation du compte et le reset MDP dépendent d'un service e-mail tiers", "Service gratuit en phase académique. En production : SendGrid / SES avec SLA"),
    ("SQLite en multi-utilisateurs", "Faible (PPE)", "SQLite non recommandé pour accès concurrent massif", "Acceptable en phase académique. Migration PostgreSQL prévue si déploiement réel"),
    ("Agrément AMF non détenu", "Faible (PPE)", "En conditions réelles, agrément requis pour conseil en investissement", "Périmètre académique défini. Conformité MiFID II démonstrable"),
]
risk_table = make_table(doc,
    ["Risque", "Probabilité", "Impact potentiel", "Mitigation"],
    risk_rows, header_color='1A2E5A')
prob_colors = {"Moyen": "FFF3CD", "Faible": "D4EDDA"}
for row in risk_table.rows[1:]:
    cell = row.cells[1]
    for kw, col in prob_colors.items():
        if kw in cell.text:
            set_cell_bg(cell, col)
            break

doc.add_paragraph()

add_heading(doc, "4.4 Recommandations pour la suite", 2)
for rec in [
    "Activer le déploiement sur Render (render.yaml déjà configuré) pour rendre l'application accessible à des testeurs externes.",
    "Migrer vers PostgreSQL si un accès multi-utilisateurs simultané est requis en production.",
    "Remplacer yfinance par une API financière professionnelle (ex. Alpha Vantage, Polygon.io) pour garantir la disponibilité des données.",
    "Compléter la couverture de tests (pytest) pour les cas limites du questionnaire, de l'optimiseur et des endpoints d'authentification.",
    "Finaliser l'intégration de scikit-learn pour des fonctionnalités de classification ML des profils utilisateurs.",
]:
    add_bullet(doc, rec)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 5. ANNEXES
# ══════════════════════════════════════════════
add_heading(doc, "5. Annexes", 1)

add_heading(doc, "Annexe A — Structure du projet (arborescence)", 2)
tree = doc.add_paragraph()
tree.style = doc.styles['No Spacing']
run = tree.add_run(
"""Projet PPE/
├── src/
│   ├── api/
│   │   ├── main.py                    ← Point d'entrée FastAPI (5 routers)
│   │   ├── endpoints.py               ← Routes principales (portfolio, profil)
│   │   ├── auth_endpoints.py          ← Auth (register, login, verify, reset)
│   │   ├── portfolio_endpoints.py     ← Cycle de vie portefeuille
│   │   ├── training_endpoints.py      ← Academy (modules, leçons, quiz)
│   │   └── gdpr_endpoints.py          ← Export et suppression données RGPD
│   ├── portfolio/
│   │   ├── optimizer.py               ← Markowitz + Black-Litterman
│   │   ├── recommender.py             ← Sélection et filtrage des actifs
│   │   ├── explainer.py               ← Explications textuelles (Core/Diversif.)
│   │   └── filters.py                 ← Filtrage avancé de l'univers
│   ├── profile/
│   │   ├── suggestion_rules.py        ← Logique Safety-First MiFID II
│   │   ├── suggestion_engine.py       ← Moteur de suggestions temps réel
│   │   ├── profile_engine.py          ← Moteur de classification des profils + cohérence MiFID II
│   │   └── profile_builder.py         ← Construction du profil refiné
│   ├── database/
│   │   └── models.py                  ← 8 modèles SQLAlchemy (+ AuthToken)
│   ├── auth/ + utils/                 ← JWT, email_service, config
│   └── data/                          ← data_loader, storage, universe
├── templates/                         ← 16 fichiers HTML
│   ├── explorer.html                  ← Explorateur univers actifs [NOUVEAU]
│   ├── reset_password.html            ← Reset MDP [NOUVEAU]
│   └── verify_status.html             ← Vérification e-mail [NOUVEAU]
├── static/
│   └── js/                            ← 7 fichiers JS (api.js, ui.js, cookie-consent...)
├── scripts/                           ← 24 scripts utilitaires
├── data/
│   └── universe.csv                   ← 479 actifs — 405 Actions + 74 ETFs (cache 48h)
├── render.yaml                        ← Configuration déploiement Render [NOUVEAU]
└── requirements.txt                   ← 30 dépendances Python (versions fixes)"""
)
run.font.name = 'Courier New'
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x1A, 0x2E, 0x5A)

doc.add_paragraph()

add_heading(doc, "Annexe B — Dépendances Python (requirements.txt)", 2)
deps_rows = [
    ("fastapi==0.129.0", "Framework API REST"),
    ("uvicorn==0.40.0", "Serveur ASGI"),
    ("yfinance==1.1.0", "Données financières (Yahoo Finance)"),
    ("pandas==3.0.0", "Manipulation de données"),
    ("numpy==2.4.2", "Calculs matriciels"),
    ("pydantic==2.12.5", "Validation des données"),
    ("sqlalchemy==2.0.46", "ORM base de données"),
    ("python-multipart==0.0.22", "Formulaires HTTP"),
    ("python-jose==3.5.0", "Tokens JWT"),
    ("passlib==1.7.4", "Hachage des mots de passe"),
    ("bcrypt==4.2.1", "Chiffrement"),
    ("python-dotenv==1.2.1", "Variables d'environnement"),
    ("requests==2.32.5", "Appels HTTP"),
    ("jinja2==3.1.6", "Moteur de templates"),
    ("pyportfolioopt==1.5.6", "Optimisation Markowitz"),
    ("scikit-learn==1.8.0", "Machine Learning (classification profils)"),
    ("scipy==1.17.0", "Solveur mathématique"),
    ("joblib==1.5.3", "Parallélisation"),
    ("plotly==5.24.1", "Visualisation interactive"),
    ("tenacity==9.1.4", "Retry automatique (résilience yfinance)"),
    ("email-validator==2.3.0", "Validation emails"),
    ("h11==0.16.0", "Protocole HTTP/1.1"),
    ("httpx==0.28.1", "Client HTTP async (tests)"),
    ("cvxpy==1.8.1", "Solveur d'optimisation convexe"),
    ("osqp==1.1.1", "Solveur OSQP (via CVXPY)"),
    ("scs==3.2.11", "Solveur SCS (via CVXPY)"),
    ("ecos==2.0.14", "Solveur ECOS (via CVXPY)"),
    ("watchfiles==1.1.1", "Rechargement à chaud (dev)"),
    ("websockets==16.0", "WebSockets"),
]

deps_table = make_table(doc, ["Package (requirements.txt)", "Usage"], deps_rows, header_color='1F6EB4')
for row in deps_table.rows[1:]:
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)

doc.add_paragraph()
doc.add_page_break()

final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final.add_run("— Fin de l'étude de faisabilité — Mars 2026 —")
r.italic = True
r.font.color.rgb = RGBColor(0x99,0x99,0x99)
r.font.size = Pt(10)

doc.save(OUTPUT_PATH)
print(f"✅ Document Word généré : {OUTPUT_PATH}")
